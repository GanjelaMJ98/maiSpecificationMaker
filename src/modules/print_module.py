from docx import Document
from docx.shared import Inches



def print_specification(project_name, system_name, table_product, sum,path = None):
    document = Document()
    document.add_heading('Specification', 0)
    p = document.add_paragraph('Project: ')
    p.add_run(project_name).bold = True
    p = document.add_paragraph('System: ')
    p.add_run(system_name).italic = True
    table = document.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'code'
    hdr_cells[1].text = 'product'
    hdr_cells[2].text = 'price'
    hdr_cells[3].text = 'sum'
    hdr_cells[4].text = 'contr'
    hdr_cells[5].text = 'fab'
    hdr_cells[6].text = 'cur'
    hdr_cells[7].text = 'del'
    for a, b, c, d, e, f, g, h in table_product:
        row_cells = table.add_row().cells
        row_cells[0].text = str(a)
        row_cells[1].text = str(b)
        row_cells[2].text = str(c)
        row_cells[3].text = str(d)
        row_cells[4].text = str(e)
        row_cells[5].text = str(f)
        row_cells[6].text = str(g)
        row_cells[7].text = str(h)
    p = document.add_paragraph('Sum: ' + str(sum))
    document.add_page_break()
    if path is not None:
        document.save(path)
    else:
        document.save('demo.docx')
    return 0